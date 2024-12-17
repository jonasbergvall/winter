import json
import streamlit as st
from pyairtable import Table
import toml
import requests
import io
from streamlit_option_menu import option_menu
from docx import Document
from docx.shared import Inches
from PIL import Image, ImageDraw, ImageOps
import os

# Extract Airtable configuration from Streamlit secrets
AIRTABLE_API_KEY = st.secrets["airtable"]["api_key"]
AIRTABLE_BASE_ID = st.secrets["airtable"]["base_id"]
AIRTABLE_TABLE_NAME = st.secrets["airtable"]["table_name"]


# Load JSON Data
with open('winter.json') as user_data_file:
    user_data = json.load(user_data_file)

with open('stories.json') as stories_file:
    stories_data = json.load(stories_file)

with open('diverse_elements.json') as diverse_file:
    diverse_data = json.load(diverse_file)

with open('antidote.json', 'r') as modes_file:
    modes_data = json.load(modes_file)

# Global Styling for Streamlit App
logo = '''
<div style="text-align: right; margin-bottom: 20px;">
    <img src="https://bestofworlds.se/img/lglogo.png" width="150px">
</div>
'''

streamlit_style = """
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Roboto+Slab:wght@400;700&display=swap');
    html, body, div, h1, h2, h3, h4, h5, h6, p, span, input, textarea, select, button, [class*="css"] {
        font-family: 'Roboto Slab', serif !important;
    }
    h3 {
        font-size: 22px !important;
        text-align: left;
    }
    h4 {
        font-size: 18px !important;
        text-align: left;
        color: #333;
    }
    p {
        font-size: 14px !important;
        line-height: 1.6;
        color: gray;
        margin-bottom: 10px;
    }
    .explanation-text {
        font-size: 12px !important;
        line-height: 1.4 !important;
        color: #666666 !important;
        margin-top: 10px;
    }
    </style>
"""

st.markdown(logo, unsafe_allow_html=True)
st.markdown(streamlit_style, unsafe_allow_html=True)

# Option Menu for Navigation (now placed above the title)
selected = option_menu(
    menu_title=None,
    options=["Start", "Introspection", "Your Story", "Analysis"],
    icons=["house", "pencil", "book", "graph-up"],
    menu_icon="menu-button-fill",
    default_index=0,
    orientation="horizontal"
)

# Variables to Store Scores
if "scores" not in st.session_state:
    st.session_state["scores"] = {
        "Leader": 0,
        "Fighter": 0,
        "Visionary": 0,
        "Lover": 0
    }

# Function to round corners of an image
def round_corners(image, radius):
    # Create mask for rounded corners
    mask = Image.new('L', image.size, 0)
    draw = ImageDraw.Draw(mask)
    draw.rounded_rectangle([(0, 0), image.size], radius=radius, fill=255)

    # Apply the mask to the image to round corners
    result = ImageOps.fit(image, mask.size, centering=(0.5, 0.5))
    result.putalpha(mask)

    return result

# Navigation Logic
if selected == "Start":
    st.title("Winter Narratives")
    st.image("https://www.bestofworlds.se/img/winter.webp", use_container_width=True)
    st.write("Welcome to the Winter Narrative app! Explore your feelings about winter and see which story fits you best.")
    st.write("This is a fun way to explore diversity and different perspectives, especially regarding something as universally experienced yet uniquely perceived as winter. Continue by clicking Introspection in the menu.")

elif selected == "Introspection":
    st.title("The essence of winter")
    st.write("This is not a test, see it as just an exploration of your perspectives. Choose an option which feels most right to you.")
    for question_key, question_data in user_data["user_input"].items():
        question = question_data["question"]
        options = [option["text"] for option in question_data["options"]]
        choice = st.radio(question, options, key=question_key)

        # Update Scores Based on User Choice
        if choice:
            for option in question_data["options"]:
                if option["text"] == choice:
                    for archetype, score in option["scores"].items():
                        st.session_state["scores"][archetype] += score

    if st.button("OK"):
        st.write("Thank you for answering the questions! You can now proceed to the 'See Your Story' section.")

elif selected == "Your Story":
    st.title("Your Winter Story")
    # Determine the User's Archetype or Blended Archetype
    max_score = max(st.session_state["scores"].values())
    high_scorers = [archetype for archetype, score in st.session_state["scores"].items() if score == max_score]

    if len(high_scorers) == 1:
        # Pure Story
        selected_archetype = high_scorers[0]
        story = stories_data["stories"]["pure"][selected_archetype]
        st.subheader(story["title"])
        st.write(story["story"])

        # Add Enriching Your Winter Experience content
        enriching_content = diverse_data["diverse_elements"][selected_archetype]
        st.markdown("### Enriching Your Winter Experience")
        enriching_text = "\n\n".join([perspective["perspective"] for perspective in enriching_content["additional_perspectives"]])
        st.write(enriching_text)

        # Add download button for generated story
        def generate_docx_from_text(story_title, story_text, enriching_text):
            """
            Generates a DOCX file with the given story title, story content, and enriching content.
            """
            doc = Document()

            # Add the logo to the document
            logo_url = "https://bestofworlds.se/img/lglogo.png"
            try:
                response = requests.get(logo_url)
                response.raise_for_status()
                logo_image = io.BytesIO(response.content)
                doc.add_picture(logo_image, width=Inches(2))
                doc.add_paragraph()  # Add a blank line after the logo
            except requests.exceptions.RequestException as e:
                st.error(f"Failed to download logo image: {e}")

            # Add story title and story text to document
            doc.add_heading(story_title, level=1)
            doc.add_paragraph(story_text)

            # Add enriching content to the document
            doc.add_heading("Enriching Your Winter Experience", level=2)
            doc.add_paragraph(enriching_text)

            # Save the document to a BytesIO buffer
            docx_output = io.BytesIO()
            doc.save(docx_output)
            docx_output.seek(0)
            return docx_output

        # Generate and download the DOCX file
        docx_file = generate_docx_from_text(story["title"], story["story"], enriching_text)
        st.download_button(
            label="Download Your Winter Narrative as DOCX",
            data=docx_file,
            file_name="winter_narrative.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    elif len(high_scorers) == 2:
        # Blended Story
        archetype_combo = f"{high_scorers[0]}_{high_scorers[1]}"
        story = stories_data["stories"]["blended"][archetype_combo]
        st.subheader(story["title"])
        st.write(story["story"])
    else:
        # In case all archetypes are tied (unlikely, but handled here)
        st.write("It looks like you equally embody all aspects of winter! Please try answering again to see if we can find your true match.")


elif selected == "Analysis":
    st.title("Diverse Winter Narratives")
    st.write("Winter isn't just a season we experience—it's a story we all tell in different ways. The way we appreciate winter reflects deeper aspects of who we are, shaped by the values and ideas within our communities. To explore this, we’ve identified 10 unique themes—key perspectives on winter that appear again and again in books and films. These themes aren’t the only possible ones, nor are they always separate. They often blend together, adding richness and complexity to how we understand winter. By looking at these different lenses, we can better appreciate the beauty of our diverse experiences and how they shape our shared story of winter.")

    # Dropdown menus and images for analysis section
    col1, col2, col3 = st.columns(3)

    # Adding "Select" option to modes list
    modes_list = ["Select"] + list(modes_data.keys())

    with col1:
        mode1 = st.selectbox("Select a mode", modes_list, key="mode1")
        if mode1 and mode1 != "Select":
            image_path = os.path.join("images", f"{mode1.lower()}.jpg")
            if os.path.exists(image_path):
                image1 = Image.open(image_path)
                image1_rounded = round_corners(image1, radius=30)
                st.image(image1_rounded, caption=mode1, use_container_width=True)
            st.markdown(f"<p class='explanation-text'>{modes_data[mode1]['Explanation']}</p>", unsafe_allow_html=True)
        elif mode1 == "Select":
            neutral_image_path = os.path.join("images", "neutral.jpg")
            if os.path.exists(neutral_image_path):
                neutral_image = Image.open(neutral_image_path)
                neutral_image_rounded = round_corners(neutral_image, radius=30)
                st.image(neutral_image_rounded, caption="Please select a mode", use_container_width=True)

    with col3:
        mode2 = st.selectbox("Select another mode", modes_list, key="mode2")
        if mode2 and mode2 != "Select":
            image_path = os.path.join("images", f"{mode2.lower()}.jpg")
            if os.path.exists(image_path):
                image2 = Image.open(image_path)
                image2_rounded = round_corners(image2, radius=30)
                st.image(image2_rounded, caption=mode2, use_container_width=True)
            st.markdown(f"<p class='explanation-text'>{modes_data[mode2]['Explanation']}</p>", unsafe_allow_html=True)
        elif mode2 == "Select":
            neutral_image_path = os.path.join("images", "neutral.jpg")
            if os.path.exists(neutral_image_path):
                neutral_image = Image.open(neutral_image_path)
                neutral_image_rounded = round_corners(neutral_image, radius=30)
                st.image(neutral_image_rounded, caption="Please select a mode", use_container_width=True)

    with col2:
        st.write("Select a theme from each dropdown menu. What happens when these two perspectives come together? Do they clash, or do they complement each other? Discover how different ways of seeing winter can highlight the tensions and harmonies in our shared experience.")

    # Add a section to gather user's thoughts about winter
    st.markdown("### Share Your Thoughts")
    st.write("Do you have further thoughts on winter? Please share!")
    user_input = st.text_area("Share your reflections:")

    # Agreement text
    st.markdown(
        "By clicking 'Share,' you agree to the terms outlined in the disclaimers, including the use of your input for research and improvement purposes."
    )
    if st.button("Share"):
        if user_input.strip():  # Ensure the input isn't empty
            try:
                # Connect to Airtable
                table = Table(AIRTABLE_API_KEY, AIRTABLE_BASE_ID, AIRTABLE_TABLE_NAME)
                
                # Add feedback to Airtable
                table.create({"Feedback": user_input})
                
                st.success("Your thoughts have been saved successfully!")
            except Exception as e:
                st.error(f"An error occurred while saving your feedback: {e}")
        else:
            st.warning("Please enter your thoughts before saving.")

# Footer with Disclaimer Text and Link without Underline
st.markdown("<hr>", unsafe_allow_html=True)
st.markdown(
    """
    <style>
    .tooltip {
        position: relative;
        display: inline-block;
        cursor: pointer;
        font-size: 14px;
        color: #007BFF; /* Blue text for interactivity */
    }
    .tooltip .tooltiptext {
        visibility: hidden;
        width: 300px;
        background-color: #555;
        color: #fff;
        text-align: center;
        border-radius: 6px;
        padding: 5px;
        position: absolute;
        z-index: 1;
        bottom: 125%; /* Position above */
        left: 50%;
        margin-left: -150px;
        opacity: 0;
        transition: opacity 0.3s;
    }
    .tooltip:hover .tooltiptext {
        visibility: visible;
        opacity: 1;
    }
    a {
        text-decoration: none; /* Remove underline from links */
        color: #007BFF; /* Match link color */
    }
    a:hover {
        text-decoration: underline; /* Optional: underline on hover */
    }
    </style>

    <div style="text-align: center; margin-top: -10px;">
        <p class="tooltip">Disclaimer
            <span class="tooltiptext">
                Disclaimer: The narratives and perspectives provided are generated for exploratory and entertainment purposes. 
                They may not reflect real-life experiences or outcomes. Please use them as a creative tool rather than factual advice.<br><br>
                Disclaimer: Your input may be stored and used for research and improvement of this application. 
                No personally identifiable information is collected or shared.<br><br>
                Note: This app uses third-party services like Airtable for data storage. 
                By submitting your thoughts, you agree to the processing of data in accordance with their privacy policies.
            </span>
        </p>
        <p>Explore more data-driven perspectives at <a href="https://bestofworlds.se" target="_blank">Best of Worlds</a>.</p>
    </div>
    """,
    unsafe_allow_html=True,
)
